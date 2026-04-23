
import io
import json
import os
import datetime as dt
from typing import Any, Dict, List

import streamlit as st

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    import docx
except Exception:
    docx = None


APP_NAME = "Write and Come Back"
APP_VERSION = "v0.11.2"

TEAM_STEPS = [
    "Welcome",
    "Load Step 0 JSON",
    "Lean Validation",
    "Approve & Export",
]

WRITER_QUICK_STEPS = [
    "Welcome",
    "Load Approved Step 0 Context",
    "Your Starting Point",
    "Write Your Core Idea",
    "Quick Writing Starter",
    "Write Now & Export",
]

WRITER_FULL_STEPS = [
    "Verify Authorities",
    "Tighten the Claim",
    "Strengthen Support",
    "Improve Writing Pack",
    "Export Refined Pack",
]

TIME_EST = {
    "team_0": "1 min",
    "team_1": "1–2 min",
    "team_2": "3–8 min",
    "team_3": "1–2 min",
    "writer_quick_0": "1–2 min",
    "writer_quick_1": "1–3 min",
    "writer_quick_2": "1–3 min",
    "writer_quick_3": "3–5 min",
    "writer_quick_4": "1–2 min",
    "writer_quick_5": "1–2 min",
    "writer_full_0": "2–4 min",
    "writer_full_1": "2–4 min",
    "writer_full_2": "2–4 min",
    "writer_full_3": "2–4 min",
    "writer_full_4": "1–2 min",
}


def init_state():
    if "project" not in st.session_state:
        st.session_state.project = {
            "version": APP_VERSION,
            "created_at": dt.datetime.now().isoformat(timespec="seconds"),
            "updated_at": dt.datetime.now().isoformat(timespec="seconds"),
            "mode": "team",  # team | writer
            "phase": "quick",  # for writer only
            "step": 0,
            "answers": {},
            "outputs": {},
            "flags": [],
            "feedback": {},
            "api_log": [],
            "step0_loaded": False,
        }


def touch():
    st.session_state.project["updated_at"] = dt.datetime.now().isoformat(timespec="seconds")


def ans(key: str, default=None):
    return st.session_state.project["answers"].get(key, default)


def set_ans(key: str, value: Any):
    st.session_state.project["answers"][key] = value
    touch()


def add_flag(text: str):
    if text and text not in st.session_state.project["flags"]:
        st.session_state.project["flags"].append(text)
        touch()


def get_openai_client():
    api_key = None
    try:
        api_key = st.secrets.get("OPENAI_API_KEY")
    except Exception:
        pass
    api_key = api_key or os.getenv("OPENAI_API_KEY")
    if not api_key or OpenAI is None:
        return None
    return OpenAI(api_key=api_key)


def call_model_json(stage: str, system_prompt: str, user_prompt: str, fallback: Dict[str, Any]) -> Dict[str, Any]:
    client = get_openai_client()
    if client is None:
        st.session_state.project["api_log"].append({
            "stage": stage, "mode": "mock", "time": dt.datetime.now().isoformat(timespec="seconds")
        })
        return fallback
    try:
        resp = client.chat.completions.create(
            model="gpt-4.1",
            temperature=0.35,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        data = json.loads(resp.choices[0].message.content or "{}")
        st.session_state.project["api_log"].append({
            "stage": stage, "mode": "openai", "time": dt.datetime.now().isoformat(timespec="seconds")
        })
        return data or fallback
    except Exception as e:
        st.warning(f"Model call failed at {stage}. Using fallback instead.")
        st.session_state.project["api_log"].append({
            "stage": stage, "mode": "fallback", "error": str(e), "time": dt.datetime.now().isoformat(timespec="seconds")
        })
        return fallback


def quote_box(title: str, body: str):
    st.markdown(f"""
    <div style="background:#f7f9fc;border:1px solid #dce3ee;border-radius:12px;padding:12px 14px;margin-bottom:12px">
      <div style="font-size:0.84rem;color:#445; font-weight:700; margin-bottom:6px">{title}</div>
      <div style="font-size:0.95rem;color:#222; line-height:1.45">{body}</div>
    </div>
    """, unsafe_allow_html=True)


def get_steps():
    mode = st.session_state.project["mode"]
    phase = st.session_state.project["phase"]
    if mode == "team":
        return TEAM_STEPS, "team"
    if phase == "quick":
        return WRITER_QUICK_STEPS, "writer_quick"
    return WRITER_FULL_STEPS, "writer_full"


def progress():
    steps, prefix = get_steps()
    step = st.session_state.project["step"]
    pct = int((step + 1) / len(steps) * 100)
    st.markdown("""
    <style>
    .stepcard{border:1px solid #dce3ee;border-radius:14px;background:#f8fbff;padding:10px;min-height:90px}
    .stepcard.active{background:#173b7a;color:white;border-color:#173b7a}
    .stepcard.done{background:#eef5ff}
    .flag{display:inline-block;background:#fff3cd;border:1px solid #ffe08a;color:#7a5b00;border-radius:999px;padding:4px 10px;font-size:.82rem;margin:3px 6px 0 0}
    .req{display:inline-block;border-radius:999px;padding:3px 9px;font-size:.75rem;background:#eef5ff;color:#173b7a;margin-right:6px}
    .opt{display:inline-block;border-radius:999px;padding:3px 9px;font-size:.75rem;background:#f5f5f5;color:#666;margin-right:6px}
    .good{display:inline-block;border-radius:999px;padding:3px 9px;font-size:.75rem;background:#ecfdf3;color:#12743b;margin-right:6px}
    .warn{display:inline-block;border-radius:999px;padding:3px 9px;font-size:.75rem;background:#fff7ed;color:#9a3412;margin-right:6px}
    </style>
    """, unsafe_allow_html=True)
    mode_label = "Team Mode" if st.session_state.project["mode"] == "team" else ("Writer Mode — Quick Pass" if st.session_state.project["phase"] == "quick" else "Writer Mode — Full Pass")
    st.write(f"**{APP_NAME}** · {APP_VERSION} · {mode_label}")
    st.progress(pct / 100)
    cols = st.columns(len(steps))
    for i, name in enumerate(steps):
        cls = "stepcard active" if i == step else ("stepcard done" if i < step else "stepcard")
        key = f"{prefix}_{i}"
        cols[i].markdown(
            f'<div class="{cls}"><div style="font-size:.73rem;font-weight:700">STEP {i+1}</div><div style="font-size:.88rem;font-weight:600;line-height:1.2;margin-top:6px">{name}</div><div style="font-size:.74rem;opacity:.85;margin-top:6px">{TIME_EST.get(key,"")}</div></div>',
            unsafe_allow_html=True,
        )


def feedback_sidebar():
    steps, _ = get_steps()
    step = st.session_state.project["step"]
    label = f"{st.session_state.project['mode']}::{st.session_state.project.get('phase','')}::{steps[step]}"
    fb = st.session_state.project["feedback"].setdefault(label, {})
    st.sidebar.markdown("## Notes for the next version")
    fb["worked"] = st.sidebar.text_area("What worked", value=fb.get("worked", ""), key=f"worked_{label}")
    fb["confusing"] = st.sidebar.text_area("What was confusing", value=fb.get("confusing", ""), key=f"conf_{label}")
    fb["too_much"] = st.sidebar.text_area("What felt too much", value=fb.get("too_much", ""), key=f"too_{label}")
    fb["missing"] = st.sidebar.text_area("What is missing", value=fb.get("missing", ""), key=f"miss_{label}")
    fb["notes"] = st.sidebar.text_area("Other notes", value=fb.get("notes", ""), key=f"notes_{label}")
    touch()


def parse_step0_json(data: Dict[str, Any]):
    set_ans("step0_raw", data)
    set_ans("step0_materials_summary", data.get("materials_summary") or data.get("materials_materials_summary") or "")
    set_ans("step0_legal_arguments", data.get("legal_arguments") or data.get("materials_legal_arguments") or [])
    set_ans("step0_authorities", data.get("authorities") or data.get("materials_authorities_mentioned") or [])
    set_ans("step0_flags", data.get("flags") or data.get("extraction_flags") or [])
    validation = data.get("validation", {})
    set_ans("step0_validation", validation)
    for f in ans("step0_flags", []):
        add_flag(f if isinstance(f, str) else str(f))
    st.session_state.project["step0_loaded"] = True
    touch()


def validated_step0_payload() -> Dict[str, Any]:
    raw = ans("step0_raw", {}) or {}
    payload = dict(raw)
    payload["legal_arguments"] = ans("validated_legal_arguments", []) or ans("step0_legal_arguments", [])
    payload["authorities"] = ans("validated_authorities", []) or ans("step0_authorities", [])
    payload["flags"] = ans("validated_flags", []) or ans("step0_flags", [])
    payload["validation"] = {
        "status": ans("validation_status", "not_ready"),
        "arguments_checked": ans("check_arguments", False),
        "quotes_checked": ans("check_quotes", False),
        "authorities_checked": ans("check_authorities", False),
        "issues_found": ans("validation_issues", ""),
        "reviewer": ans("validation_reviewer", ""),
        "date": ans("validation_date", ""),
    }
    return payload


def make_machine_export() -> str:
    return json.dumps(st.session_state.project, indent=2, ensure_ascii=False)


def make_validated_step0_json() -> str:
    return json.dumps(validated_step0_payload(), indent=2, ensure_ascii=False)


def make_human_export() -> bytes:
    if docx is None:
        return make_machine_export().encode("utf-8")
    d = docx.Document()
    p = st.session_state.project
    d.add_heading(f"{APP_NAME} — Pre-Writing Kit", 0)
    d.add_paragraph(f"Version: {p['version']}")
    d.add_paragraph(f"Created: {p['created_at']}")
    d.add_paragraph(f"Updated: {p['updated_at']}")
    d.add_heading("Step 0 Context", level=1)
    if ans("step0_materials_summary"):
        d.add_paragraph(str(ans("step0_materials_summary")))
    if ans("step0_legal_arguments"):
        d.add_paragraph("Extracted legal arguments:")
        for item in ans("step0_legal_arguments", []):
            if isinstance(item, dict):
                d.add_paragraph(str(item.get("argument", item)), style="List Bullet")
            else:
                d.add_paragraph(str(item), style="List Bullet")
    d.add_heading("Quick Pass", level=1)
    for key in ["quick_selected_point_user_text", "quick_core_idea", "quick_result_summary"]:
        if ans(key):
            d.add_paragraph(str(ans(key)))
    if ans("quick_outline"):
        d.add_heading("Quick Outline", level=2)
        for item in ans("quick_outline", []):
            d.add_paragraph(str(item), style="List Number")
    if ans("quick_writing_package"):
        wp = ans("quick_writing_package")
        d.add_heading("Quick Writing Package", level=2)
        d.add_paragraph("Abstract seed:")
        d.add_paragraph(str(wp.get("abstract_seed", "")))
        d.add_paragraph("Opening options:")
        for item in wp.get("opening_options", []):
            d.add_paragraph(str(item), style="List Bullet")
        d.add_paragraph("Section cues:")
        for k, v in wp.get("section_cues", {}).items():
            d.add_paragraph(f"{k}: {v}", style="List Bullet")
        d.add_paragraph("Write-now instruction:")
        d.add_paragraph(str(wp.get("write_now_prompt", "")))
    if ans("full_refined_claim") or ans("full_support_summary") or ans("full_improved_writing_summary"):
        d.add_heading("Full Pass", level=1)
        if ans("full_refined_claim"):
            d.add_paragraph(f"Refined claim: {ans('full_refined_claim')}")
        if ans("full_support_summary"):
            d.add_paragraph(f"Support summary: {ans('full_support_summary')}")
        if ans("full_improved_writing_summary"):
            d.add_paragraph(f"Writing pack improvements: {ans('full_improved_writing_summary')}")
    if p.get("flags"):
        d.add_heading("Carry-forward Flags", level=1)
        for f in p["flags"]:
            d.add_paragraph(str(f), style="List Bullet")
    bio = io.BytesIO()
    d.save(bio)
    bio.seek(0)
    return bio.getvalue()


def make_audit_report() -> bytes:
    if docx is None:
        return make_validated_step0_json().encode("utf-8")
    d = docx.Document()
    d.add_heading("Step 0.5 Audit Report", 0)
    d.add_paragraph(f"Reviewer: {ans('validation_reviewer', '')}")
    d.add_paragraph(f"Date: {ans('validation_date', '')}")
    d.add_paragraph(f"Status: {ans('validation_status', 'not_ready')}")
    d.add_heading("What was checked", level=1)
    d.add_paragraph(f"Arguments checked: {ans('check_arguments', False)}")
    d.add_paragraph(f"Quotes checked: {ans('check_quotes', False)}")
    d.add_paragraph(f"Authorities checked: {ans('check_authorities', False)}")
    d.add_heading("Issues Found", level=1)
    d.add_paragraph(ans("validation_issues", "") or "None recorded.")
    d.add_heading("Argument Review", level=1)
    for i, item in enumerate(ans("validated_legal_arguments", []) or ans("step0_legal_arguments", []), 1):
        d.add_paragraph(f"Argument {i}", style=None)
        if isinstance(item, dict):
            d.add_paragraph(str(item.get("argument", "")))
            d.add_paragraph(f"Reviewer action: {item.get('review_action', '')}")
            d.add_paragraph(f"Reviewer note: {item.get('review_note', '')}")
        else:
            d.add_paragraph(str(item))
    d.add_heading("Authority Review", level=1)
    for i, item in enumerate(ans("validated_authorities", []) or ans("step0_authorities", []), 1):
        if isinstance(item, dict):
            d.add_paragraph(f"{i}. {item.get('name', '')}")
            d.add_paragraph(f"Status: {item.get('team_status', '')}")
            d.add_paragraph(f"Reviewer note: {item.get('team_note', '')}")
        else:
            d.add_paragraph(str(item))
    bio = io.BytesIO()
    d.save(bio)
    bio.seek(0)
    return bio.getvalue()


def clear_downstream(mode: str, from_step: int):
    if mode == "team":
        mapping = {
            1: ["step0_raw", "step0_materials_summary", "step0_legal_arguments", "step0_authorities", "step0_flags"],
            2: ["validated_legal_arguments", "validated_authorities", "validated_flags", "check_arguments", "check_quotes", "check_authorities", "validation_issues", "validation_status", "validation_reviewer", "validation_date"],
        }
    elif mode == "writer_quick":
        mapping = {
            1: ["step0_raw", "step0_materials_summary", "step0_legal_arguments", "step0_authorities"],
            2: ["quick_selected_point_source", "quick_selected_point_user_text", "quick_result_summary", "quick_outline", "quick_working_title", "quick_working_question"],
            3: ["quick_core_idea", "quick_goal_signal", "quick_weak_point"],
            4: ["quick_writing_package"],
        }
    else:
        mapping = {
            0: ["full_verified_authorities", "full_authority_additions", "full_authority_removals"],
            1: ["full_refined_claim"],
            2: ["full_support_summary", "full_support_needs"],
            3: ["full_improved_writing_pack", "full_improved_writing_summary"],
        }
    for idx, keys in mapping.items():
        if idx >= from_step:
            for key in keys:
                st.session_state.project["answers"].pop(key, None)
                st.session_state.project["outputs"].pop(key, None)
    touch()


def revisit_warning(mode: str, from_step: int):
    st.warning("You may revisit an earlier step. But if you change a response and run the model again, later outputs should be rebuilt from this step forward.")
    if st.button("Clear downstream outputs from this step"):
        clear_downstream(mode, from_step)
        st.success("Later outputs were cleared. Re-run from this step.")
        st.rerun()


def ensure_writer_has_approved_step0():
    validation = ans("step0_validation", {})
    if not validation or validation.get("status") != "approved":
        st.error("This Step 0 file is not yet approved for writer use.")
        st.stop()


st.set_page_config(page_title=APP_NAME, layout="wide")
init_state()
progress()

with st.sidebar:
    st.markdown(f"### {APP_NAME}")
    st.caption(f"{APP_VERSION} · {'OpenAI connected' if get_openai_client() else 'Mock mode'}")
    mode_choice = st.radio("Workspace", ["Team Mode", "Writer Mode"], index=0 if st.session_state.project["mode"] == "team" else 1)
    st.session_state.project["mode"] = "team" if mode_choice == "Team Mode" else "writer"
    if st.session_state.project["mode"] == "writer":
        phase_choice = st.radio("Writer Phase", ["Quick Pass", "Full Pass"], index=0 if st.session_state.project["phase"] == "quick" else 1)
        st.session_state.project["phase"] = "quick" if phase_choice == "Quick Pass" else "full"
    st.divider()
    feedback_sidebar()
    st.divider()
    st.download_button("Export workspace checkpoint (.json)", data=make_machine_export(), file_name=f"workspace_{APP_VERSION}.json", mime="application/json", use_container_width=True)
    uploaded_cp = st.file_uploader("Load saved workspace checkpoint", type=["json"], key="load_checkpoint")
    if uploaded_cp and st.button("Load this checkpoint", use_container_width=True):
        st.session_state.project = json.load(uploaded_cp)
        st.success("Checkpoint loaded.")
        st.rerun()

mode = st.session_state.project["mode"]
phase = st.session_state.project["phase"]
step = st.session_state.project["step"]

# TEAM MODE
if mode == "team":
    if step == 0:
        st.markdown("## Team Mode — Step 1 — Welcome")
        quote_box(
            "What this mode does",
            "Team Mode is for Step 0.5 only. Step 0 is done externally in ChatGPT or Claude. This validator is intentionally lean: it checks whether the output is reliable and verifiable enough before the writer sees it."
        )
        quote_box(
            "What this mode does not do",
            "This is not a deep editorial rewriting stage. The goal is not to perfect the output. The goal is to catch what would break trust."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["team_0"]}</span>', unsafe_allow_html=True)
        if st.button("Start Validation", type="primary"):
            st.session_state.project["step"] = 1
            touch()
            st.rerun()

    elif step == 1:
        st.markdown("## Team Mode — Step 2 — Load Step 0 JSON")
        revisit_warning("team", 1)
        quote_box(
            "What to do here",
            "Upload the Step 0 JSON generated externally. The validator will parse it, display the extracted arguments and authorities, and prepare a lean verification workflow."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["team_1"]}</span>', unsafe_allow_html=True)
        step0_file = st.file_uploader("Upload Step 0 JSON", type=["json"], key="team_step0_json")
        if step0_file and st.button("Load Step 0 for validation", type="primary"):
            data = json.load(step0_file)
            parse_step0_json(data)
            st.session_state.project["step"] = 2
            st.rerun()

        if st.session_state.project["step0_loaded"]:
            st.markdown("### Step 0 summary")
            st.write(ans("step0_materials_summary"))
            st.markdown("**Extracted legal arguments**")
            for i, item in enumerate(ans("step0_legal_arguments", []), 1):
                text = item.get("argument") if isinstance(item, dict) else str(item)
                st.markdown(f"{i}. {text}")

    elif step == 2:
        st.markdown("## Team Mode — Step 3 — Lean Validation")
        revisit_warning("team", 2)
        quote_box(
            "What to do here",
            "Use this as a fast trust check. Confirm what looks accurate, flag what looks wrong, and make only light corrections if necessary. If this step takes too long, it is becoming too heavy."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["team_2"]}</span>', unsafe_allow_html=True)

        st.markdown("### Minimum checks")
        check_arguments = st.checkbox("Arguments are faithful to the materials", value=ans("check_arguments", False))
        check_quotes = st.checkbox("Quotes are traceable to the cited source/location", value=ans("check_quotes", False))
        check_authorities = st.checkbox("Authorities are sane enough for Step 1", value=ans("check_authorities", False))

        st.markdown("### Argument review")
        validated_arguments: List[Any] = []
        for i, item in enumerate(ans("step0_legal_arguments", []), 1):
            if isinstance(item, dict):
                arg = dict(item)
            else:
                arg = {"argument": str(item)}
            st.markdown(f"**Argument {i}**")
            arg_text = st.text_area(f"Argument text {i}", value=arg.get("argument", ""), height=90, key=f"arg_text_{i}")
            quote = arg.get("supporting_quote", "") or arg.get("quote", "")
            if quote:
                st.caption(f'Supporting quote: "{quote}"')
            if arg.get("source") or arg.get("location"):
                st.caption(f"Source: {arg.get('source', '')} | Location: {arg.get('location', '')}")
            review_action = st.radio(f"Reviewer action {i}", ["Confirm", "Revise", "Remove"], index=0, horizontal=True, key=f"arg_action_{i}")
            review_note = st.text_input(f"Reviewer note {i} (optional)", value=arg.get("review_note", ""), key=f"arg_note_{i}")
            arg["argument"] = arg_text
            arg["review_action"] = review_action
            arg["review_note"] = review_note
            if review_action != "Remove":
                validated_arguments.append(arg)

        st.markdown("### Authority review")
        validated_authorities: List[Any] = []
        for i, item in enumerate(ans("step0_authorities", []), 1):
            if isinstance(item, dict):
                a = dict(item)
            else:
                a = {"name": str(item)}
            name = st.text_input(f"Authority {i}", value=a.get("name", ""), key=f"auth_name_{i}")
            link = st.text_input(f"Link {i}", value=a.get("verified_link") or a.get("link") or "", key=f"auth_link_{i}")
            team_status = st.radio(f"Authority status {i}", ["Looks correct", "Needs check", "Incorrect"], index=0, horizontal=True, key=f"auth_status_{i}")
            team_note = st.text_input(f"Authority note {i} (optional)", value=a.get("team_note", ""), key=f"auth_note_{i}")
            a["name"] = name
            if link:
                a["verified_link"] = link
            a["team_status"] = team_status
            a["team_note"] = team_note
            if team_status != "Incorrect":
                validated_authorities.append(a)

        st.markdown("### One critical question")
        validation_issues = st.text_area(
            "Is there anything here that is clearly wrong or misleading?",
            value=ans("validation_issues", ""),
            height=100,
        )

        if st.button("Save lean validation", type="primary"):
            set_ans("validated_legal_arguments", validated_arguments)
            set_ans("validated_authorities", validated_authorities)
            set_ans("validated_flags", ans("step0_flags", []))
            set_ans("check_arguments", check_arguments)
            set_ans("check_quotes", check_quotes)
            set_ans("check_authorities", check_authorities)
            set_ans("validation_issues", validation_issues)
            st.session_state.project["step"] = 3
            st.rerun()

    elif step == 3:
        st.markdown("## Team Mode — Step 4 — Approve & Export")
        quote_box(
            "What to do here",
            "Set the final status. Only approved Step 0 files should be handed to the writer for Step 1 ingestion."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["team_3"]}</span>', unsafe_allow_html=True)

        validation_status = st.radio("Validation status", ["not_ready", "needs_fix", "approved"], index=["not_ready", "needs_fix", "approved"].index(ans("validation_status", "not_ready")), horizontal=True)
        reviewer = st.text_input("Reviewer name", value=ans("validation_reviewer", ""))
        date_str = st.text_input("Date", value=ans("validation_date", dt.date.today().isoformat()))

        set_ans("validation_status", validation_status)
        set_ans("validation_reviewer", reviewer)
        set_ans("validation_date", date_str)

        if validation_status == "approved":
            st.success("This Step 0 file is approved for writer use.")
        elif validation_status == "needs_fix":
            st.warning("This Step 0 file needs fixes before writer use.")
        else:
            st.info("This Step 0 file is not ready yet.")

        st.markdown("### Exports")
        st.download_button(
            "Download validated Step 0 JSON",
            data=make_validated_step0_json(),
            file_name="validated_step0.json",
            mime="application/json",
            type="primary",
        )
        st.download_button(
            "Download Step 0.5 audit report (.docx)",
            data=make_audit_report(),
            file_name="step0_5_audit_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# WRITER MODE
else:
    if step == 0:
        st.markdown("## Step 1 — Welcome")
        quote_box(
            "What this platform does",
            "It takes the lecture context already prepared and approved by your team, then helps you move quickly from that context to a workable writing starter. Quick Pass is lean on purpose. The goal is to help you write before you edit."
        )
        quote_box(
            "What to expect",
            "Quick Pass comes first. It should feel light and forward-moving. Full Pass comes only after Quick Pass is done, and only if you want to make the paper stronger."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["writer_quick_0"] if phase=="quick" else TIME_EST["writer_full_0"]}</span>', unsafe_allow_html=True)
        if st.button("Start Writer Flow", type="primary"):
            st.session_state.project["step"] = 1 if phase == "quick" else 0
            touch()
            st.rerun()

    elif phase == "quick" and step == 1:
        st.markdown("## Step 2 — Load Approved Step 0 Context")
        revisit_warning("writer_quick", 1)
        quote_box(
            "What to do here",
            "Upload the approved Step 0 JSON. The app will ingest the lecture context, show what was already extracted, and let you confirm your own starting point without redoing the preparation work."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["writer_quick_1"]}</span>', unsafe_allow_html=True)
        step0_file = st.file_uploader("Upload approved Step 0 JSON", type=["json"], key="writer_step0_json")
        if step0_file and st.button("Load approved Step 0 context", type="primary"):
            data = json.load(step0_file)
            parse_step0_json(data)
            ensure_writer_has_approved_step0()
            # override visible items with validated ones where present
            val_args = data.get("legal_arguments")
            val_auth = data.get("authorities")
            if val_args is not None:
                set_ans("step0_legal_arguments", val_args)
            if val_auth is not None:
                set_ans("step0_authorities", val_auth)
            st.session_state.project["step"] = 2
            st.rerun()
        if st.session_state.project["step0_loaded"]:
            ensure_writer_has_approved_step0()
            st.markdown("### Approved Step 0 summary")
            st.write(ans("step0_materials_summary"))

    elif phase == "quick" and step == 2:
        st.markdown("## Step 3 — Your Starting Point")
        revisit_warning("writer_quick", 2)
        quote_box(
            "What to do here",
            "From your lecture, these are the legal arguments already identified. You do not need to choose perfectly. Pick one, combine them, or write your own version. This is just your starting point."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["writer_quick_2"]}</span>', unsafe_allow_html=True)
        st.markdown("### Extracted legal arguments")
        args = ans("step0_legal_arguments", [])
        for i, item in enumerate(args, 1):
            st.markdown(f"{i}. {item.get('argument') if isinstance(item, dict) else item}")
        source_choice = st.text_area("What do you want to write about?", value=ans("quick_selected_point_user_text", ""), height=130, placeholder="Write 1–2 sentences in your own words. You may pick one of the arguments, combine them, or write your own version.")
        additional = st.text_area("Optional: Is there anything else you want the app to keep in mind before we continue?", value=ans("quick_extra_consideration", ""), height=90)
        if st.button("Continue from this starting point", type="primary"):
            set_ans("quick_selected_point_source", "writer-defined")
            set_ans("quick_selected_point_user_text", source_choice.strip())
            set_ans("quick_extra_consideration", additional)
            st.session_state.project["step"] = 3
            st.rerun()

    elif phase == "quick" and step == 3:
        st.markdown("## Step 4 — Write Your Core Idea")
        revisit_warning("writer_quick", 3)
        quote_box(
            "What to do here",
            "This is the most important step in Quick Pass. In 3–5 sentences, explain your paper idea in your own words. Do not worry about structure, grammar, or completeness. Just explain what you want to say."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["writer_quick_3"]}</span>', unsafe_allow_html=True)
        core_idea = st.text_area("Write your core idea", value=ans("quick_core_idea", ""), height=180, placeholder="Write in your own words. Keep moving. Do not edit too much.")
        st.markdown("### Optional quick signals")
        goal_signal = st.text_area("What result or conclusion do you want the paper to reach?", value=ans("quick_goal_signal", ""), height=90)
        weak_point = st.text_area("What part of your idea feels weakest right now?", value=ans("quick_weak_point", ""), height=90)
        if st.button("Build the quick writing starter", type="primary"):
            set_ans("quick_core_idea", core_idea)
            set_ans("quick_goal_signal", goal_signal)
            set_ans("quick_weak_point", weak_point)
            fallback = {
                "result_summary": "A workable writing starter has been built from Step 0 context and the writer’s own explanation.",
                "working_title": "Working paper title",
                "working_question": "What should this paper prove or clarify based on the writer’s core idea?",
                "outline": ["Introduction and legal problem", "Legal framework or doctrine", "Main legal argument", "What follows from that argument", "Conclusion"],
            }
            system = """
You are helping a writer move fast. Use Step 0 context and the writer's own words.
Return JSON with:
result_summary
working_title
working_question
outline
Do not invent a new argument. Push forward what is already there.
"""
            user = f"""Step 0 summary: {ans("step0_materials_summary")}
Writer-defined starting point: {ans("quick_selected_point_user_text")}
Writer's core idea: {core_idea}
Desired result: {goal_signal}
Weak point: {weak_point}
Extra consideration: {ans("quick_extra_consideration","")}
"""
            with st.spinner("Building the quick starter..."):
                data = call_model_json("quick_core_idea", system, user, fallback)
            set_ans("quick_result_summary", data.get("result_summary"))
            set_ans("quick_working_title", data.get("working_title"))
            set_ans("quick_working_question", data.get("working_question"))
            set_ans("quick_outline", data.get("outline", []))
            st.session_state.project["step"] = 4
            st.rerun()

    elif phase == "quick" and step == 4:
        st.markdown("## Step 5 — Quick Writing Starter")
        revisit_warning("writer_quick", 4)
        quote_box(
            "What to do here",
            "This step gives you a rough but usable writing starter: an abstract seed, opening options, and section cues. The point is not to perfect the paper. The point is to help you start writing now."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["writer_quick_4"]}</span>', unsafe_allow_html=True)
        if st.button("Create quick writing starter", type="primary"):
            outline = ans("quick_outline", [])
            fallback = {
                "abstract_seed": "This article examines a legal issue surfaced in the lecture and clarifies the legal position it takes.",
                "opening_options": ["Start with the legal problem directly.", "Start with the doctrinal tension already visible in the lecture.", "Start with the consequence your paper wants to show."],
                "section_cues": {x: f"What do you want to say in '{x}' in your own words?" for x in outline},
                "write_now_prompt": "Write one section now in your own words. Do not edit yet. Keep going.",
            }
            system = """
Create a rough writing starter from the existing context.
Return JSON with:
abstract_seed
opening_options
section_cues
write_now_prompt
This is Quick Pass. Keep it light and useful.
"""
            user = f"""Step 0 context summary: {ans("step0_materials_summary")}
Writer starting point: {ans("quick_selected_point_user_text")}
Writer's core idea: {ans("quick_core_idea")}
Current outline: {ans("quick_outline", [])}
"""
            with st.spinner("Creating the quick writing starter..."):
                data = call_model_json("quick_writing_starter", system, user, fallback)
            set_ans("quick_writing_package", data)
            st.session_state.project["step"] = 5
            st.rerun()

    elif phase == "quick" and step == 5:
        st.markdown("## Step 6 — Write Now & Export")
        quote_box(
            "What to do here",
            "This is the end of Quick Pass. You now have enough to start writing. Do not edit first. Write first. Export your files, then draft one section in your own words."
        )
        st.markdown(f'<span class="req">Required</span><span class="opt">Time: {TIME_EST["writer_quick_5"]}</span>', unsafe_allow_html=True)
        wp = ans("quick_writing_package") or {}
        if wp:
            st.info(wp.get("write_now_prompt", "Write one section now in your own words. Do not edit yet."))
        st.download_button("Download human-readable pre-writing kit (.docx)", data=make_human_export(), file_name=f"pre_writing_kit_{APP_VERSION}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
        st.download_button("Download machine-readable checkpoint (.json)", data=make_machine_export(), file_name=f"checkpoint_{APP_VERSION}.json", mime="application/json")
        if st.button("Move to Full Pass"):
            st.session_state.project["phase"] = "full"
            st.session_state.project["step"] = 0
            touch()
            st.rerun()

    elif phase == "full" and step == 0:
        st.markdown("## Full Pass — Step 1 — Verify Authorities")
        revisit_warning("writer_full", 0)
        quote_box(
            "What to do here",
            "Full Pass begins only after Quick Pass. This first step shows the authorities identified in Step 0. Confirm what belongs, remove what does not, and add anything missing."
        )
        st.markdown(f'<span class="req">Required for Full Pass</span><span class="opt">Time: {TIME_EST["writer_full_0"]}</span>', unsafe_allow_html=True)
        authorities = ans("step0_authorities", [])
        confirmed, removals = [], []
        additions = st.text_area("Add any missing authority, case, law, or rule", value=ans("full_authority_additions", ""), height=90)
        for i, item in enumerate(authorities, 1):
            label = item.get("name") if isinstance(item, dict) else str(item)
            link = item.get("verified_link") if isinstance(item, dict) else None
            keep = st.checkbox(f"Keep {i}. {label}" + (f" — {link}" if link else ""), value=True, key=f"auth_keep_{i}")
            if keep:
                confirmed.append(item)
            else:
                removals.append(item)
        if st.button("Confirm authorities", type="primary"):
            set_ans("full_verified_authorities", confirmed)
            set_ans("full_authority_removals", removals)
            set_ans("full_authority_additions", additions)
            st.session_state.project["step"] = 1
            st.rerun()

    elif phase == "full" and step == 1:
        st.markdown("## Full Pass — Step 2 — Tighten the Claim")
        revisit_warning("writer_full", 1)
        quote_box("What to do here", "Now that you already have a starter, rewrite the one claim your paper must prove. Keep it in one sentence.")
        st.markdown(f'<span class="req">Required for Full Pass</span><span class="opt">Time: {TIME_EST["writer_full_1"]}</span>', unsafe_allow_html=True)
        refined_claim = st.text_area("What is the one claim your paper must prove?", value=ans("full_refined_claim", ""), height=100)
        if st.button("Tighten claim", type="primary"):
            set_ans("full_refined_claim", refined_claim)
            st.session_state.project["step"] = 2
            st.rerun()

    elif phase == "full" and step == 2:
        st.markdown("## Full Pass — Step 3 — Strengthen Support")
        revisit_warning("writer_full", 2)
        quote_box("What to do here", "This step helps identify where your argument still needs stronger support. Use it to name the weakest supported part of the paper.")
        st.markdown(f'<span class="req">Required for Full Pass</span><span class="opt">Time: {TIME_EST["writer_full_2"]}</span>', unsafe_allow_html=True)
        support_need = st.text_area("What part of the paper still needs stronger support?", value=ans("full_support_needs", ""), height=110)
        if st.button("Strengthen support", type="primary"):
            set_ans("full_support_needs", support_need)
            fallback = {"support_summary": "The paper needs stronger support in the area identified by the writer.", "new_flags": []}
            system = """
Strengthen support for an existing paper starter.
Return JSON with:
support_summary
new_flags
"""
            user = f"""Refined claim: {ans("full_refined_claim")}
Support need: {support_need}
Verified authorities: {ans("full_verified_authorities", [])}
Added authorities: {ans("full_authority_additions", "")}
"""
            with st.spinner("Strengthening support..."):
                data = call_model_json("full_strengthen_support", system, user, fallback)
            set_ans("full_support_summary", data.get("support_summary"))
            for f in data.get("new_flags", []):
                add_flag(f)
            st.session_state.project["step"] = 3
            st.rerun()

    elif phase == "full" and step == 3:
        st.markdown("## Full Pass — Step 4 — Improve Writing Pack")
        revisit_warning("writer_full", 3)
        quote_box("What to do here", "This step improves the writing package you already have. It should make it easier to keep writing, not restart the whole process.")
        st.markdown(f'<span class="req">Required for Full Pass</span><span class="opt">Time: {TIME_EST["writer_full_3"]}</span>', unsafe_allow_html=True)
        improve_note = st.text_area("What part of the writing package do you want improved most?", value=ans("full_improve_note", ""), height=100)
        if st.button("Improve writing pack", type="primary"):
            set_ans("full_improve_note", improve_note)
            fallback = {
                "writing_summary": "The writing package has been improved for clearer drafting.",
                "improved_writing_pack": {
                    "extra_opening_options": ["Start with the sharpest legal conflict first."],
                    "extra_section_prompt": "In your own words, explain why this section matters to the paper’s main claim."
                }
            }
            system = """
Improve an existing writing pack after quick pass.
Return JSON with:
writing_summary
improved_writing_pack
"""
            user = f"""Current quick writing package: {ans("quick_writing_package")}
Refined claim: {ans("full_refined_claim")}
Support summary: {ans("full_support_summary")}
Improvement requested: {improve_note}
"""
            with st.spinner("Improving writing pack..."):
                data = call_model_json("full_improve_writing_pack", system, user, fallback)
            set_ans("full_improved_writing_summary", data.get("writing_summary"))
            set_ans("full_improved_writing_pack", data.get("improved_writing_pack"))
            st.session_state.project["step"] = 4
            st.rerun()

    elif phase == "full" and step == 4:
        st.markdown("## Full Pass — Step 5 — Export Refined Pack")
        quote_box("What to do here", "Full Pass is complete. Export the updated files and continue writing. The point of this stage was to strengthen what already existed, not to stop momentum.")
        st.markdown(f'<span class="req">Required for Full Pass</span><span class="opt">Time: {TIME_EST["writer_full_4"]}</span>', unsafe_allow_html=True)
        st.download_button("Download updated human-readable pre-writing kit (.docx)", data=make_human_export(), file_name=f"pre_writing_kit_{APP_VERSION}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
        st.download_button("Download updated machine-readable checkpoint (.json)", data=make_machine_export(), file_name=f"checkpoint_{APP_VERSION}.json", mime="application/json")
